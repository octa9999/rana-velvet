from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rana-Velvet-Admin-Panel-Guide.docx"
LOGIN = ROOT / "docs" / "qa-artifacts" / "admin-sign-in.png"
CATALOG = ROOT / "docs" / "qa-artifacts" / "storefront-catalog.png"

INK = RGBColor(28, 31, 29)
GREEN = RGBColor(13, 107, 63)
MUTED = RGBColor(95, 100, 96)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        shade(cell, "0D6B3F")
        cell_text(cell, header, bold=True, color=RGBColor(255, 255, 255))
        set_cell_margins(cell)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            cell_text(cells[index], value)
            set_cell_margins(cells[index])
    doc.add_paragraph()
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in (("Title", 28, INK), ("Heading 1", 17, GREEN), ("Heading 2", 12.5, INK), ("Heading 3", 11, GREEN)):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Title"
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(11)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)
    header = section.header.paragraphs[0]
    header.text = "RANA VELVET  |  ADMIN PANEL GUIDE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Rana Velvet  |  Client operating guide")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def title(doc, text):
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def note(doc, heading, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "EEF5F1")
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    a = p.add_run(heading + " ")
    a.bold = True
    a.font.color.rgb = GREEN
    p.add_run(text)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


doc = Document()
style_doc(doc)

title(doc, "Rana Velvet Admin Panel Guide")
p = doc.add_paragraph("A simple operating guide for products, orders, customer requests, site content and daily checks.")
p.runs[0].font.size = Pt(13)
p.runs[0].font.color.rgb = MUTED
doc.add_paragraph("Version: August 2026", style="Subtitle")
note(doc, "Purpose.", "Use this guide to make routine changes safely. Keep product, pricing and customer data accurate, and use only approved staff accounts.")
doc.add_heading("Quick Start", level=1)
numbered(doc, [
    "Open the website admin sign-in page: /login.",
    "Enter the authorized admin email and password. Do not share the password in WhatsApp, email or screenshots.",
    "After sign-in, use the left menu to open the relevant workspace.",
    "Make one clear change at a time, save it, then confirm the public website or admin list reflects the update.",
])
if LOGIN.exists():
    doc.add_picture(str(LOGIN), width=Inches(6.45))
    c = doc.add_paragraph("Figure 1. Admin sign-in. 1) Use the approved email address. 2) Enter the password. 3) Select Sign In to open the dashboard.")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(8.5)

doc.add_page_break()
doc.add_heading("Dashboard And Navigation", level=1)
doc.add_paragraph("The dashboard is the control centre. The left menu groups work into catalog, operations and site content.")
add_table(doc, ["Area", "Use it for"], [
    ("Products / Categories / Inventory", "Manage what customers can buy, how it is grouped and available stock."),
    ("Orders / Print Orders / Checkout / Payments", "Review sales, prepare fulfilment, control checkout information and record payment progress."),
    ("Appointments / Inquiries", "Handle consultations, measurement visits, customer messages and requests."),
    ("Banners / Testimonials / Media", "Update visible home-page content and reusable image files."),
    ("Reports / Settings", "Review operational data and maintain store settings."),
], [2.0, 4.45])
note(doc, "Daily routine.", "Start with Orders, Inquiries and Appointments. Then check Inventory before publishing any product as active.")
doc.add_heading("Finding Records", level=2)
bullets(doc, [
    "Use the search field in a workspace to find a product name, SKU, customer detail, order number or QA/reference label.",
    "Use status filters to focus on active, draft, archived, open or completed records.",
    "Use Export CSV before a large review or handover. Keep exported customer files private.",
    "Import CSV only with a reviewed template and a small test file first. Check the import result before importing a full catalog.",
])

doc.add_heading("Products", level=1)
doc.add_paragraph("Products control product detail pages, collection listings, search, price and stock.")
doc.add_heading("Add A Product", level=2)
numbered(doc, [
    "Open Products and choose Add Product.",
    "Enter a clear product name, unique SKU and web-safe slug. Example slug: velvet-lounge-chair.",
    "Select the correct category, status, price, stock, material, colour and image.",
    "Choose Draft while the product is being prepared. Choose Active only after all public details are ready.",
    "Save Product. Search the new product in the list, then open its public product page to check the listing, price, photo and availability.",
])
doc.add_heading("Edit, Archive Or Delete", level=2)
bullets(doc, [
    "Edit: use the pencil action, make the change and save. Recheck the public product page.",
    "Archive: use Archived for products that should remain in records but should not be sold.",
    "Delete: delete only duplicate, test or permanently retired records. Confirm the delete prompt carefully; deletion cannot be undone.",
    "Images: add approved product images in Media or through the product image upload control. Remove unused or incorrect media after confirming no page needs it.",
])

doc.add_page_break()
doc.add_heading("Categories, Images And Inventory", level=1)
doc.add_heading("Categories", level=2)
numbered(doc, [
    "Open Categories and select Add Category.",
    "Set the name, unique slug, image and order. Use a parent category when creating a subcategory.",
    "Save, then check the storefront menu and product filters.",
    "Do not delete a category while products still rely on it. Move or update those products first.",
])
doc.add_heading("Media", level=2)
bullets(doc, [
    "Use Media to upload reusable banners and product imagery.",
    "Give files meaningful names and use optimized, correctly cropped images.",
    "Before removing an image, confirm it is not used by a product, banner or testimonial.",
])
doc.add_heading("Inventory", level=2)
numbered(doc, [
    "Open Inventory and locate the SKU or product name.",
    "Adjust stock only after counting or receiving confirmation from the workshop/showroom.",
    "Check available stock, reserved stock and the public stock state after saving.",
    "Use a short internal note/reference where available, especially for manual corrections.",
])
note(doc, "Stock rule.", "Never publish an active product with uncertain stock. Review stock after every order cancellation, return or manual adjustment.")

doc.add_heading("Orders, Checkout And Payments", level=1)
doc.add_paragraph("Orders are created by checkout and should be reviewed in operational sequence.")
add_table(doc, ["Workspace", "What to do"], [
    ("Orders", "Open new orders, confirm customer details, items, delivery information and status."),
    ("Print Orders", "Prepare the fulfilment or workshop printout once the order is checked."),
    ("Payments", "Record or verify payment status and method according to the approved payment process."),
    ("Checkout", "Maintain checkout text, options and rules. Test a change using a non-customer test order before publishing widely."),
    ("Inventory", "Confirm the sold quantity and reserved/available stock remain correct after order changes."),
], [1.55, 4.9])
doc.add_heading("Order Handling Steps", level=2)
numbered(doc, [
    "Open the order and check the order number, customer contact details, products, quantities and delivery address.",
    "Verify payment status before starting work or dispatching.",
    "Prepare the Print Order when the order is ready for the workshop or fulfilment team.",
    "Update the order status as work progresses. Keep customer-facing status accurate.",
    "For cancellation or correction, review payment and inventory impact before changing the order.",
])

doc.add_page_break()
doc.add_heading("Bookings And Customer Requests", level=1)
doc.add_paragraph("Public forms arrive in the following admin queues. Search by customer name, email, subject or reference label.")
add_table(doc, ["Public page/form", "Admin destination", "How to handle"], [
    ("Contact", "Inquiries", "Reply, assign follow-up and keep the message status current."),
    ("Custom Furniture", "Inquiries > Custom Furniture", "Review dimensions, preferences and budget before contacting the customer."),
    ("Customize Curtain: request", "Inquiries > Curtain Requests", "Check measurements, fabric, accessories and reference detail."),
    ("Partners", "Inquiries > Partner Applications", "Review company/contact details and progress the partnership response."),
    ("Consultation", "Appointments", "Confirm preferred date, time and consultation type."),
    ("Curtain measurement visit", "Appointments > Measurement Visits", "Confirm address, visit date/time and team availability."),
    ("Checkout", "Orders, Print Orders, Payments, Inventory and notifications", "Process as an order using the operational sequence in this guide."),
], [2.05, 2.2, 2.2])
doc.add_heading("Appointments", level=2)
bullets(doc, [
    "Use Appointments for consultations and Measurement Visits for curtain visits.",
    "Confirm date, time, customer contact details and address before changing status to confirmed.",
    "Reschedule only after agreeing a replacement slot. Record the final agreed time.",
])
doc.add_heading("Inquiries", level=2)
bullets(doc, [
    "Keep every request in the correct queue using the filtered pages listed above.",
    "Do not delete genuine customer messages. Mark them resolved or archive only when appropriate.",
    "Use clear internal notes and follow-up actions so another team member can continue the conversation.",
])

doc.add_heading("Site Content, Reports And Settings", level=1)
add_table(doc, ["Workspace", "Safe use"], [
    ("Banners", "Update hero/promotional content. Check desktop and mobile appearance after saving."),
    ("Testimonials", "Add approved customer feedback only. Check wording, customer consent and display order."),
    ("Reports", "Review sales, orders, inquiries and operating trends. Export only when needed and store securely."),
    ("Settings", "Maintain approved business contact details, shipping/payment settings and store configuration. Change one setting at a time and test the related public page."),
], [1.6, 4.85])
if CATALOG.exists():
    doc.add_picture(str(CATALOG), width=Inches(6.45))
    c = doc.add_paragraph("Figure 2. Storefront catalog check. After publishing or editing a product, confirm the public listing, filter behaviour, image, title and price on desktop and mobile.")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(8.5)

doc.add_page_break()
doc.add_heading("Safe Operating Rules", level=1)
bullets(doc, [
    "Use a named admin account. Do not share logins or leave the dashboard open on a shared computer.",
    "Create a clearly labelled QA record for testing and delete it immediately after the full test is complete.",
    "Never edit, delete or export a customer/order record unless you are certain it is the intended record.",
    "Use Draft before Active for products, banners and major content changes.",
    "Check mobile and desktop after changes to navigation, banners, product images or long text.",
    "Before a bulk import, export a backup and test one row first.",
    "For delete prompts, pause and verify the record name, SKU, order number or customer reference.",
])
doc.add_heading("Common Troubleshooting", level=1)
add_table(doc, ["Problem", "Check first", "Next step"], [
    ("Product is not visible", "Status, stock, category, slug and image", "Save again, refresh the public products page and check search/filter selection."),
    ("Product save is rejected", "Required fields, unique SKU/slug and valid category", "Keep the editor open, correct the error and retry. Do not create duplicate records."),
    ("Form cannot be found", "Customer email, subject and reference", "Search Inquiries or Appointments, then open the relevant filtered queue."),
    ("Order looks incorrect", "Items, address, payment and inventory", "Pause fulfilment, correct the record through the appropriate workspace and recheck stock."),
    ("Image looks wrong", "Media file, crop and product/banner assignment", "Replace the image, save and check desktop/mobile again."),
    ("Cannot sign in", "Approved email, password and account access", "Use the authorized account owner or administrator. Do not repeatedly guess passwords."),
], [1.55, 2.25, 2.65])
note(doc, "Support handover.", "When reporting an issue, include the page URL, record name or order number, time of the issue, a screenshot and the exact action that failed.")

doc.add_heading("Before You Finish", level=1)
numbered(doc, [
    "Confirm the save message or status update in the admin panel.",
    "Refresh the relevant admin list and make sure the record is still correct.",
    "Check the public storefront when the change affects customers.",
    "Remove only temporary QA/test records created by you.",
    "Log out on shared devices.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Rana Velvet Admin Panel Guide"
doc.core_properties.author = "Rana Velvet"
doc.save(OUT)
print(OUT)
